
import io
from datetime import datetime

from django.db.models import Count, Q


def generate_tickets_xlsx(tickets):
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Отчёт по тикетам'

    # Стили заголовка
    header_font = Font(name='Arial', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='2B5797', end_color='2B5797', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )

    # Заголовки столбцов
    headers = [
        '№ п/п', '№ Тикета', 'Тема обращения', 'Категория',
        'Приоритет', 'Статус', 'Автор', 'Исполнитель',
        'Дата создания', 'Дата закрытия',
    ]

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Данные
    data_alignment = Alignment(vertical='center', wrap_text=True)
    for row_idx, ticket in enumerate(tickets, 2):
        row_data = [
            row_idx - 1,
            ticket.pk,
            ticket.title,
            str(ticket.category),
            str(ticket.priority),
            str(ticket.status),
            str(ticket.author),
            str(ticket.assignee) if ticket.assignee else '—',
            ticket.created_at.strftime('%d.%m.%Y %H:%M'),
            ticket.closed_at.strftime('%d.%m.%Y %H:%M') if ticket.closed_at else '—',
        ]
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = data_alignment
            cell.border = thin_border

    # Автоширина столбцов
    for col_idx in range(1, len(headers) + 1):
        max_length = max(
            len(str(ws.cell(row=row, column=col_idx).value or ''))
            for row in range(1, ws.max_row + 1)
        )
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_length + 4, 40)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def generate_tickets_docx(tickets, title='Отчёт по тикетам'):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Заголовок
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Дата формирования
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_run.font.size = Pt(9)
    date_run.italic = True

    doc.add_paragraph()  # Отступ

    # Сводка
    total = tickets.count()
    doc.add_paragraph(f'Всего тикетов в отчёте: {total}')

    if total == 0:
        doc.add_paragraph('Нет данных для отображения.')
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # Таблица
    headers = ['№', 'Тема', 'Категория', 'Приоритет', 'Статус', 'Автор', 'Дата создания']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки таблицы
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        header_cells[idx].text = header_text
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Данные таблицы
    for num, ticket in enumerate(tickets, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = ticket.title[:60]
        row_cells[2].text = str(ticket.category)
        row_cells[3].text = str(ticket.priority)
        row_cells[4].text = str(ticket.status)
        row_cells[5].text = str(ticket.author)
        row_cells[6].text = ticket.created_at.strftime('%d.%m.%Y')

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
